"""Generic Markdown resume → Word (.docx) converter for the Resume Optimizer workflow.
Usage: python scripts/create_docx.py output/filename.md
"""
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0, 70, 127)
GREY = RGBColor(80, 80, 80)


def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _runs(p, text, size=10, default_color=None):
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.font.bold = True
        else:
            r = p.add_run(part)
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        if default_color and not r.font.bold:
            r.font.color.rgb = default_color


def add_name(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.font.bold = True


def add_contact(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.color.rgb = GREY


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text.upper())
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = BLUE
    _hr(doc)


def add_job_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.bold = True


def add_meta_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _runs(p, text, size=10, default_color=GREY)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    _runs(p, text, size=10)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    _runs(p, text, size=10)


def convert(md_path: str):
    lines = Path(md_path).read_text(encoding='utf-8').splitlines()
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    name_done = False
    contact_done = False

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        if not name_done:
            add_name(doc, line)
            name_done = True
            continue

        if not contact_done and not line.startswith('#'):
            add_contact(doc, line)
            contact_done = True
            continue

        contact_done = True

        if line.startswith('## '):
            add_section_heading(doc, line[3:])
        elif line.startswith('### '):
            add_job_title(doc, line[4:])
        elif line.startswith('- '):
            add_bullet(doc, line[2:])
        elif re.search(r'\*\*[^*]+\*\*', line) and '|' in line:
            add_meta_line(doc, line)
        elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            add_job_title(doc, line[2:-2])
        else:
            add_body(doc, line)

    out = md_path.replace('.md', '.docx')
    doc.save(out)
    print(f'Saved: {out}')
    return out


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python scripts/create_docx.py output/filename.md')
        sys.exit(1)
    convert(sys.argv[1])
